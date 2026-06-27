"""
DOCX Generator — python-docx.
Genereaza document Word editabil din report_sections.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DISCLAIMER = (
    "Acest raport a fost generat automat folosind exclusiv date disponibile public "
    "din surse verificabile. Acuratetea datelor depinde de corectitudinea informatiilor "
    "din registrele publice accesate. Roland Intelligence System nu isi asuma "
    "responsabilitatea pentru decizii bazate exclusiv pe acest raport fara verificare "
    "independenta."
)


def _setup_styles(doc: Document):
    """Configureaza stilurile documentului."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(40, 40, 40)

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(99, 102, 241)
        h.font.bold = True


def _fmt_docx_num(v) -> str:
    if isinstance(v, int | float):
        return f"{v:,.0f}"
    return str(v) if v is not None else "-"


def _docx_names(items) -> list[str]:
    res = []
    for it in items or []:
        if isinstance(it, dict):
            res.append(str(it.get("nume") or it.get("name") or it.get("denumire") or it))
        else:
            res.append(str(it))
    return res


def _add_rich_fields_docx(doc, verified_data: dict):
    """Render previously-dropped rich fields into the DOCX: predictive scores,
    benchmark, actionariat/relations, AEGRM guarantees, historical OSINT, funding."""
    pred = verified_data.get("predictive_scores", {})
    if isinstance(pred, dict) and pred.get("summary"):
        doc.add_page_break()
        doc.add_heading("Scoruri Predictive Faliment", level=1)
        altman = pred.get("altman_z", {}) or {}
        piotroski = pred.get("piotroski_f", {}) or {}
        beneish = pred.get("beneish_m", {}) or {}
        zmijewski = pred.get("zmijewski_x", {}) or {}
        zmi_state = "Distres" if zmijewski.get("distress") else ("OK" if zmijewski.get("available") else "Indisponibil")
        for ln_ in [
            f"Altman Z'': {altman.get('z_score', 'N/A')} ({altman.get('zone', 'INDISPONIBIL')})",
            f"Piotroski F: {piotroski.get('f_score', 'N/A')}/{piotroski.get('max_possible', 9)} ({piotroski.get('grade', 'N/A')})",
            f"Beneish M: {beneish.get('m_score', 'N/A')} ({beneish.get('risk', 'INDISPONIBIL')})",
            f"Zmijewski X: {zmijewski.get('x_score', 'N/A')} ({zmi_state})",
        ]:
            doc.add_paragraph(ln_, style="List Bullet")
        cp = doc.add_paragraph()
        cp.add_run(f"Concluzie: {pred.get('summary', '')} ({pred.get('distress_signals', 0)} semnale)").bold = True

    bench = verified_data.get("benchmark", {})
    if isinstance(bench, dict) and bench.get("available") and bench.get("comparisons"):
        doc.add_page_break()
        doc.add_heading(f"Benchmark Sector CAEN {bench.get('caen_code', '')}", level=1)
        section_name = bench.get("caen_section_name", "")
        if section_name:
            doc.add_paragraph(f"{section_name} — {bench.get('nr_firme_sector') or '?'} firme in sector")
        try:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Indicator", "Firma", "Media sector", "Pozitie"
            for c in bench["comparisons"]:
                row = table.add_row().cells
                row[0].text = str(c.get("metric", ""))
                row[1].text = _fmt_docx_num(c.get("firma"))
                row[2].text = _fmt_docx_num(c.get("media_sector"))
                row[3].text = str(c.get("pozitie", ""))
        except Exception:
            for c in bench["comparisons"]:
                doc.add_paragraph(str(c.get("text", "")), style="List Bullet")

    act = verified_data.get("actionariat", {})
    rel = verified_data.get("relations", {})
    act_ok = isinstance(act, dict) and act.get("available")
    rel_flags = rel.get("flags", []) if isinstance(rel, dict) else []
    if act_ok or rel_flags:
        doc.add_page_break()
        doc.add_heading("Actionariat si Relatii", level=1)
        if act_ok:
            cap = act.get("capital_social")
            stare = act.get("stare", "")
            if cap or stare:
                extra = f" | Stare: {stare}" if stare else ""
                doc.add_paragraph(f"Capital social: {_fmt_docx_num(cap)}{extra}")
            for label, items in (("Asociati", act.get("asociati")), ("Administratori", act.get("administratori"))):
                names = _docx_names(items)
                if names:
                    doc.add_heading(label, level=2)
                    for n in names:
                        doc.add_paragraph(n, style="List Bullet")
        for fl in rel_flags:
            doc.add_paragraph(f"[{fl.get('severity', 'INFO')}] {fl.get('type', '')}: {fl.get('detail', '')}", style="List Bullet")

    risk = verified_data.get("risk", {})
    aegrm_field = risk.get("aegrm_guarantees", {}) if isinstance(risk, dict) else {}
    aegrm = aegrm_field.get("value") if isinstance(aegrm_field, dict) else None
    hist = verified_data.get("historical_flags", [])
    aegrm_ok = isinstance(aegrm, dict) and aegrm.get("has_data")
    hist_ok = isinstance(hist, list) and bool(hist)
    if aegrm_ok or hist_ok:
        doc.add_page_break()
        doc.add_heading("Garantii si Istoric (OSINT)", level=1)
        if aegrm_ok:
            doc.add_paragraph(f"Garantii reale mobiliare (AEGRM): {aegrm.get('count', 0)}")
            guarantees = aegrm.get("guarantees") or aegrm.get("results") or []
            if isinstance(guarantees, list):
                for g in guarantees[:8]:
                    txt = (g.get("descriere") or g.get("creditor") or g.get("title") or str(g)) if isinstance(g, dict) else str(g)
                    doc.add_paragraph(str(txt)[:200], style="List Bullet")
        if hist_ok:
            for fl in hist:
                if isinstance(fl, dict):
                    # osint_client emits {type(slug), label(human), severity, snippet};
                    # prefer the human label + snippet, fall back to other shapes.
                    label = fl.get("label") or fl.get("type") or fl.get("title") or "Semnal"
                    detail = fl.get("snippet") or fl.get("detail") or fl.get("description") or ""
                    date_raw = fl.get("date") or fl.get("data") or ""
                    doc.add_paragraph(f"{label} {date_raw}: {detail}"[:240], style="List Bullet")
                else:
                    doc.add_paragraph(str(fl), style="List Bullet")

    funding = verified_data.get("funding_programs", {})
    if isinstance(funding, dict) and funding.get("eligible"):
        doc.add_page_break()
        doc.add_heading("Programe de Finantare Eligibile", level=1)
        if funding.get("summary"):
            doc.add_paragraph(str(funding["summary"]))
        try:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Program", "Suma max (EUR)", "Termen"
            for p in funding["eligible"]:
                suma = p.get("suma_max_eur", 0)
                suma_str = f"{suma:,.0f}" if isinstance(suma, int | float) and suma else "-"
                row = table.add_row().cells
                row[0].text = str(p.get("nume", ""))
                row[1].text = suma_str
                row[2].text = str(p.get("termen", "") or "-")
        except Exception:
            for p in funding["eligible"]:
                doc.add_paragraph(f"{p.get('nume', '')} — {p.get('suma_max_eur', 0)} EUR", style="List Bullet")


def generate_docx(report_sections: dict, meta: dict, output_path: str, verified_data: dict = None):
    """Genereaza DOCX din report_sections. B15: due_diligence + early_warnings."""
    verified_data = verified_data or {}
    doc = Document()
    _setup_styles(doc)

    # Core properties (metadata)
    doc.core_properties.title = meta.get("title", "Raport RIS")
    doc.core_properties.author = "Roland Intelligence System"
    doc.core_properties.keywords = f"RIS, {meta.get('company_name', '')}, business intelligence"
    doc.core_properties.subject = meta.get("analysis_type", "Analiza")

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(meta.get("title", "Raport"))
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(99, 102, 241)
    title_run.bold = True

    company = meta.get("company_name", "")
    if company:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(company)
        cr.font.size = Pt(16)
        cr.font.color.rgb = RGBColor(80, 80, 80)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"Nivel: {meta.get('report_level', 'N/A')} | "
        f"Generat: {meta.get('generated_at', '')} | "
        f"Surse: {meta.get('sources_count', 0)}"
    )
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(130, 130, 130)

    risk = meta.get("risk_score", "N/A")
    if risk != "N/A":
        # risk_score poate fi dict {"score": 72, "label": "Verde"} sau string
        if isinstance(risk, dict):
            risk_label = risk.get("label", "N/A")
            risk_display = f"{risk.get('score', '')} ({risk_label})" if risk.get("score") else risk_label
        else:
            risk_label = str(risk)
            risk_display = risk_label
        risk_para = doc.add_paragraph()
        risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        risk_run = risk_para.add_run(f"Scor Risc: {risk_display}")
        risk_run.font.size = Pt(16)
        risk_run.bold = True
        color_map = {"Verde": RGBColor(34, 197, 94), "Galben": RGBColor(200, 150, 0), "Rosu": RGBColor(220, 50, 50)}
        risk_run.font.color.rgb = color_map.get(risk_label, RGBColor(100, 100, 100))

    doc.add_page_break()

    # 9D: Table of Contents
    doc.add_heading("Cuprins", level=1)
    # Insert Word TOC field (auto-updates on open in Word)
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr_text)
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)
    # Placeholder text — Word replaces on update
    placeholder_run = toc_para.add_run("(Apasati Ctrl+A apoi F9 pentru a actualiza cuprinsul)")
    placeholder_run.font.color.rgb = RGBColor(150, 150, 150)
    placeholder_run.font.size = Pt(9)
    placeholder_run.font.italic = True
    fld_char_end_run = toc_para.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    fld_char_end_run._r.append(fld_char_end)

    doc.add_page_break()

    # Sections
    for key, section in report_sections.items():
        title = section.get("title", key)
        content = section.get("content", "")

        doc.add_heading(title, level=1)

        for paragraph in content.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            if paragraph.startswith("**") or paragraph.startswith("##"):
                clean = paragraph.replace("**", "").replace("##", "").strip()
                doc.add_heading(clean, level=2)
            elif paragraph.startswith("- ") or paragraph.startswith("* "):
                doc.add_paragraph(paragraph[2:], style="List Bullet")
            else:
                p = doc.add_paragraph(paragraph)
                # Highlight trust labels
                if "[OFICIAL]" in paragraph:
                    for run in p.runs:
                        if "[OFICIAL]" in run.text:
                            run.font.color.rgb = RGBColor(0, 170, 0)
                elif "[ESTIMAT]" in paragraph:
                    for run in p.runs:
                        if "[ESTIMAT]" in run.text:
                            run.font.color.rgb = RGBColor(255, 136, 0)

        doc.add_paragraph()  # spacer

    # B15: Due Diligence Checklist from verified_data
    due_diligence = verified_data.get("due_diligence", {})
    dd_checklist = []
    if isinstance(due_diligence, dict):
        dd_checklist = due_diligence.get("checklist", [])
    elif isinstance(due_diligence, list):
        dd_checklist = due_diligence
    if dd_checklist:
        doc.add_page_break()
        doc.add_heading("Due Diligence Checklist", level=1)
        for item in dd_checklist[:15]:
            if isinstance(item, dict):
                name = str(item.get("name", "N/A"))
                status = item.get("status", "N/A")
                icon = "DA" if status in ("DA", True) else "NU" if status in ("NU", False) else "N/A"
                p = doc.add_paragraph()
                run = p.add_run(f"[{icon}] ")
                run.bold = True
                if icon == "DA":
                    run.font.color.rgb = RGBColor(34, 197, 94)
                elif icon == "NU":
                    run.font.color.rgb = RGBColor(220, 50, 50)
                else:
                    run.font.color.rgb = RGBColor(150, 150, 150)
                p.add_run(name)

    # B15: Early Warning Signals from verified_data
    early_warnings = verified_data.get("early_warnings", [])
    if isinstance(early_warnings, list) and early_warnings:
        doc.add_page_break()
        doc.add_heading("Semnale de Alarma (Early Warnings)", level=1)
        for ew in early_warnings[:10]:
            if isinstance(ew, dict):
                signal = str(ew.get("signal", ew.get("message", "N/A")))
                severity = ew.get("severity", "MEDIUM")
                p = doc.add_paragraph()
                sev_run = p.add_run(f"[{severity}] ")
                sev_run.bold = True
                if severity == "HIGH":
                    sev_run.font.color.rgb = RGBColor(220, 50, 50)
                elif severity == "MEDIUM":
                    sev_run.font.color.rgb = RGBColor(200, 150, 0)
                p.add_run(signal)
                detail = ew.get("detail", "")
                if detail:
                    dp = doc.add_paragraph(str(detail))
                    dp.runs[0].font.size = Pt(9)
                    dp.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            elif isinstance(ew, str):
                doc.add_paragraph(ew, style="List Bullet")

    # Rich fields previously dropped: predictive, benchmark, actionariat, AEGRM, historical, funding
    _add_rich_fields_docx(doc, verified_data)

    # Sources
    sources = meta.get("sources", [])
    if sources:
        doc.add_page_break()
        doc.add_heading("Surse Utilizate", level=1)
        for src in sources:
            if isinstance(src, dict):
                level = src.get("level", "?")
                name = src.get("name", "N/A")
                status = src.get("status", "N/A")
                doc.add_paragraph(f"[Nivel {level}] {name} - {status}", style="List Bullet")
            else:
                doc.add_paragraph(str(src), style="List Bullet")

    # Disclaimer
    doc.add_page_break()
    doc.add_heading("Disclaimer", level=2)
    disc_para = doc.add_paragraph(DISCLAIMER)
    disc_para.runs[0].font.size = Pt(8)
    disc_para.runs[0].font.italic = True
    disc_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
