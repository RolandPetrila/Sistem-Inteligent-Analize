"""
Teste pentru docx_generator — generare document Word din report_sections.
"""
import os
import tempfile


def _make_sections() -> dict:
    return {
        "executive_summary": {"title": "Rezumat Executiv", "content": "Firma este stabila."},
        "financial_analysis": {"title": "Analiza Financiara", "content": "Cifra de afaceri a crescut."},
    }


def _make_meta() -> dict:
    return {
        "title": "Raport RIS",
        "company_name": "Test SRL",
        "report_level": 2,
        "generated_at": "2026-04-08T10:00:00",
        "sources_count": 3,
        "risk_score": {"score": 72, "label": "Verde"},
        "sources": ["ANAF", "ONRC"],
    }


class TestGenerateDocx:
    def test_creeaza_fisier_docx(self):
        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_functioneaza_cu_sectiuni_goale(self):
        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx({}, _make_meta(), path)
            assert os.path.exists(path)
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_functioneaza_cu_verified_data_due_diligence(self):
        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "due_diligence": [
                {"check": "Firma activa", "status": "DA"},
                {"check": "Platitor TVA", "status": "DA"},
            ],
            "early_warnings": [
                {"warning": "Scadere CA > 30%"},
            ],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            assert os.path.exists(path)
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_sanctions_hit_rendered(self):
        """Screening sanctiuni HIT cu diacritice — nu trebuie sa arunce, fisier valid."""
        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "sanctions": {
                "status": "hit",
                "hits": [{"query": "Ștefan Popescu", "matched_name": "POPESCU, Ștefán",
                          "source": "OFAC", "type": "individual"}],
                "checked": ["Firma Țăndărei SRL", "Ștefan Popescu"],
                "lists_checked": ["OFAC", "EU", "UN"],
                "data_date": "2026-07-11T00:00:00Z", "total_entries": 53000,
            }
        }
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_eurostat_sector_rendered(self):
        """Benchmark sector UE (Eurostat) — nu trebuie sa arunce, fisier valid."""
        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "eurostat_sector": {
                "available": True, "nace_used": "J62", "nace_label": "Computer programming",
                "year": "2024",
                "indicators": {
                    "ENT_NR": {"label": "Numar firme", "ro": 45240, "eu": 1008501, "nace": "J62"},
                    "EMP_ENT_NR": {"label": "Angajati / firma", "ro": 4, "eu": 5, "nace": "J62"},
                },
            }
        }
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_seap_procurement_history_rendered(self):
        """Istoric achizitii publice SEAP cu diacritice — nu trebuie sa arunce."""
        from backend.reports.docx_generator import generate_docx

        verified_data = {"market": {"seap": {"value": {
            "contracts_verified": True,
            "total_contracts": 2, "contracts_count": 1, "direct_count": 1, "total_value": 900000,
            "contracts": [{"title": "Lucrări reabilitare școală", "value": 850000, "currency": "RON",
                           "authority": "Primăria Târgoviște", "date": "2025-01-15"}],
            "direct_acquisitions": [{"title": "Achiziție consumabile", "value": 5000,
                                     "authority": "Spitalul Județean", "date": "2024-09-10"}],
        }}}}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_tender_opportunities_rendered(self):
        """Angle A: oportunitati SICAP cu diacritice — nu trebuie sa arunce."""
        from backend.reports.docx_generator import generate_docx

        verified_data = {"tender_opportunities": {"available": True, "count": 1, "days_back": 30,
            "opportunities": [{"title": "Construcție grădiniță", "authority": "Primăria Târgoviște",
                               "cpv": "45214100-1", "value": 750000, "deadline": "2026-08-01", "notice_no": "CN5"}]}}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_functioneaza_cu_meta_incomplet(self):
        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), {"title": "Minimal"}, path)
            assert os.path.exists(path)
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_risk_factors_rendered_with_diacritics_and_severity_order(self):
        """BUG2 (2026-07-16): risk_score['factors'] was never rendered in DOCX — the
        score's actual drivers (BPI insolvency, litigation, Monitorul Oficial cesiuni)
        were invisible in this format. Fixture is 100% synthetic (repo public) and
        includes diacritics; assert CONTENT (not just 'does not raise'), and assert
        CRITICAL factors are ordered before HIGH regardless of input list order."""
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "risk_score": {
                "numeric_score": 15,
                "score": "Rosu",
                "factors": [
                    ("Firma inactiva la ANAF", "HIGH"),
                    ("ZOMBIE: CA=0 + angajati=0 + status activ - firma nu opereaza", "CRITICAL"),
                    ("Firma in insolvență - dosar deschis la Tribunalul Târgoviște", "CRITICAL"),
                    ("Litigii găsite - cauțiune și garanție reală mobiliară", "MEDIUM"),
                    ("Dosare judecătorești multiple", "LOW"),
                ],
            }
        }
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            full_text = "\n".join(paragraphs)

            assert "Factori de Risc" in full_text
            assert "[CRITICAL]" in full_text
            assert "[HIGH]" in full_text
            assert "ZOMBIE" in full_text
            assert "Tribunalul Târgoviște" in full_text
            assert "Dosare judecătorești multiple" in full_text

            # Severity ordering: both CRITICAL paragraphs before the HIGH paragraph,
            # regardless of the order they appear in the input list.
            idx_zombie = next(i for i, t in enumerate(paragraphs) if "ZOMBIE" in t)
            idx_insolventa = next(i for i, t in enumerate(paragraphs) if "Tribunalul Târgoviște" in t)
            idx_high = next(i for i, t in enumerate(paragraphs) if t.startswith("[HIGH]"))
            assert idx_zombie < idx_high
            assert idx_insolventa < idx_high
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_docx_content_all_rich_sections_rendered(self):
        """DRY #3 (2026-07-14): verifica CONTINUTUL randat (nu doar 'nu arunca'),
        pe fixture-ul populat -- companion la TestRichFields (html) + TestRichFieldsPdf."""
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "predictive_scores": {
                "altman_z": {"z_score": 2.9, "zone": "SAFE"},
                "piotroski_f": {"f_score": 7, "max_possible": 9, "grade": "STRONG"},
                "beneish_m": {"m_score": -2.5, "risk": "OK"},
                "zmijewski_x": {"x_score": -1.2, "distress": False, "available": True},
                "distress_signals": 0,
                "summary": "Indicatori in zona normala",
            },
            "benchmark": {
                "available": True, "caen_code": "6201", "caen_section_name": "IT",
                "nr_firme_sector": 1200,
                "comparisons": [{"metric": "Cifra de afaceri", "firma": 500000,
                                 "media_sector": 450000, "ratio": 1.1, "pozitie": "Peste medie"}],
            },
            "actionariat": {"available": True, "asociati": [{"nume": "Ion Popescu"}],
                            "administratori": ["Maria Ionescu"], "capital_social": 200, "stare": "activa"},
            "relations": {"flags": [{"type": "ONE_PERSON", "detail": "Admin = asociat", "severity": "INFO"}]},
            "risk": {"aegrm_guarantees": {"value": {"has_data": True, "count": 1,
                     "has_guarantees": True, "details": [
                         {"nr_inregistrare": "2024-000123", "data": "2024-03-11",
                          "creditor": "Banca Exemplu SA", "tip_bun": "Gaj auto", "status": "ACTIV"},
                     ]}}},
            "historical_flags": [{"type": "cesiune_parti_sociale",
                                  "label": "Cesiune parti sociale detectata", "severity": "HIGH",
                                  "snippet": "cesiune 60% parti sociale catre o terta persoana"}],
            "funding_programs": {"eligible": [{"nume": "Start-Up Nation", "suma_max_eur": 200000,
                                "termen": "2026-12-31", "link": "https://example.ro"}],
                                "count": 1, "summary": "1 program eligibil"},
        }
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            full_text = "\n".join(text_parts)

            assert "Scoruri Predictive Faliment" in full_text
            assert "Altman" in full_text
            assert "Benchmark Sector CAEN" in full_text
            assert "Actionariat si Relatii" in full_text
            assert "Ion Popescu" in full_text
            assert "Gaj auto" in full_text
            # osint_client shape {type, label, severity, snippet}: human label + snippet
            # must be rendered, not the raw slug (regression guard, bug fixat 2026-06-27).
            assert "Cesiune parti sociale detectata" in full_text
            assert "cesiune 60% parti sociale catre o terta persoana" in full_text
            assert "cesiune_parti_sociale" not in full_text
            assert "Start-Up Nation" in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_functioneaza_cu_rich_fields_aegrm_historical(self):
        """TASK 2: populated AEGRM + historical OSINT flags (cu diacritice) prin DOCX.
        Sectiunea 'Garantii si Istoric (OSINT)' nu e atinsa de firme curate (fara
        semnale), deci asta e singura acoperire pentru calea populata DOCX."""
        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "risk": {
                "aegrm_guarantees": {
                    "value": {
                        "has_data": True,
                        "count": 2,
                        "has_guarantees": True,
                        "details": [
                            {"nr_inregistrare": "2024-000456", "data": "2024-06-02",
                             "creditor": "Exemplu Leasing SA", "tip_bun": "Autovehicul", "status": "ACTIV"},
                            {"nr_inregistrare": "2023-000789", "data": "2023-11-20",
                             "creditor": "Banca Transilvania S.A.", "tip_bun": "Ipoteca mobiliara", "status": "RADIAT"},
                        ],
                    }
                }
            },
            # Exact shape emitted by osint_client: {type, label, severity, snippet}.
            "historical_flags": [
                {
                    "type": "cesiune_parti_sociale",
                    "label": "Cesiune părți sociale detectată",
                    "severity": "HIGH",
                    "snippet": "Schimbare asociați — cesiune 60% părți sociale",
                },
                {
                    "type": "dizolvare_lichidare",
                    "label": "Dizolvare / Lichidare / Radiere",
                    "severity": "CRITICAL",
                    "snippet": "Mențiune privind dizolvarea voluntară înregistrată la ONRC",
                },
            ],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            assert os.path.exists(path)
            assert os.path.getsize(path) > 0
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_web_intelligence_content_rendered_with_diacritics_and_dedup(self):
        """verified["web_intelligence"] (Brave Search + Jina enrichment, real quota
        spent on EVERY analysis) was rendered NOWHERE in DOCX before this fix (grep
        in backend/reports/ = 0 hits). Assert actual CONTENT (not just 'does not
        raise'), on a fixture matching the real DB shape confirmed in
        data/ris.db reports.full_data (a duplicate title+url entry observed live)."""
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {"web_intelligence": {"categories": {
            "stiri": [
                {"title": "Compania își extinde activitatea în Târgoviște",
                 "url": "https://exemplu-stiri.ro/articol", "sentiment": "positive"},
                {"title": "Compania își extinde activitatea în Târgoviște",
                 "url": "https://exemplu-stiri.ro/articol", "sentiment": "positive"},
            ],
            "juridic": [
                {"title": "Litigiu comercial înregistrat pentru societate",
                 "url": "https://exemplu-just.ro/dosar", "sentiment": "negative"},
            ],
            "recenzii": [],
        }}}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)

            assert "Prezenta Online (OSINT)" in full_text
            assert "Compania își extinde activitatea în Târgoviște" in full_text
            assert "[POZITIV]" in full_text
            assert "Litigiu comercial înregistrat pentru societate" in full_text
            assert "[NEGATIV]" in full_text
            # Dedup: identical stiri entry appears exactly once.
            assert full_text.count("Compania își extinde activitatea în Târgoviște") == 1
            # Categoria goala ("recenzii") nu apare deloc.
            assert "Recenzii" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_web_intelligence_absent_omits_section(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, {})
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Prezenta Online" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)


class TestFinancialRatiosDocx:
    """2026-07-16 ("RIS colecteaza > afiseaza"): HTML has _build_financial_ratios_html,
    PDF has a dedicated E6 section, DOCX had 0 code (grep "financial_ratios" in
    docx_generator.py = 0 hits before this fix). Real values from data/ris.db
    (job 85ec7fff, TAROM CUI 477647) -- repo public."""

    def test_real_tarom_ratios_rendered_as_table(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {"risk_score": {"financial_ratios": [
            {"name": "Marja Profit Net", "value": 23.39, "unit": "%", "interpretation": "Excelent"},
            {"name": "ROA", "value": 25.19, "unit": "%", "interpretation": "Excelent"},
            {"name": "Rata Capitalizare", "value": -9.23, "unit": "%", "interpretation": "Subcapitalizat"},
            {"name": "CA per Angajat", "value": 1130414, "unit": "RON", "interpretation": ""},
        ]}}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            full_text = "\n".join(text_parts)

            assert "Indicatori Financiari" in full_text
            assert "Marja Profit Net" in full_text
            assert "23.39%" in full_text
            assert "ROA" in full_text
            assert "25.19%" in full_text
            assert "1,130,414 RON" in full_text
            assert "Excelent" in full_text
            assert "Subcapitalizat" in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_empty_ratios_omits_section(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, {"risk_score": {"financial_ratios": []}})
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Indicatori Financiari" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_absent_key_omits_section(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, {})
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Indicatori Financiari" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)


class TestMapsRatingKeyTakeawaysSectorPositionDocx:
    """2026-07-16 ("RIS colecteaza > afiseaza", etajul 3): 3 campuri calculate corect,
    randate in 0/8 formate inainte de acest fix. Fixture-urile folosesc formele si
    valorile REALE gasite in data/ris.db (job 85ec7fff, TAROM CUI 477647 -- repo
    public)."""

    def test_all_three_rendered_with_real_shapes(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {
            "maps_rating": {
                "found": True, "name": "TAROM", "rating": 3.3, "reviews_count": 767,
                "address": "Calea Bucurestilor 224F, 075100 Otopeni", "source": "google_maps",
            },
            "key_takeaways": (
                "• Cu o cifra de afaceri de 1,226,498,739 RON, TAROM prezinta o baza "
                "financiara solida pentru parteneriat.\n"
                "• Capitalurile proprii negative de -105,192,156 RON indica un risc de "
                "insolventa tehnica ce necesita monitorizare."
            ),
            "benchmark": {"available": True, "caen_code": "5110", "comparisons": [
                {"metric": "Cifra de afaceri", "firma": 1226498739, "media_sector": 3300000000, "ratio": 0.37, "pozitie": "Sub medie"},
            ]},
            "risk_score": {"sector_position": {
                "Cifra de afaceri": {"ratio_vs_avg": 0.37, "estimated_percentile": "sub P25"},
            }},
        }
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            full_text = "\n".join(text_parts)

            assert "Puncte Cheie" in full_text
            assert "financiara solida pentru parteneriat" in full_text
            assert "Prezenta pe Google Maps" in full_text
            assert "3.3/5" in full_text
            assert "767 recenzii" in full_text
            assert "Pozitie in Sector" in full_text
            assert "sub P25" in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_maps_rating_not_found_omits_section(self):
        """Real shape: {"found": False, "error": "no_results", "source":
        "google_maps"} -- must NOT render '0 stele'."""
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {"maps_rating": {"found": False, "error": "no_results", "source": "google_maps"}}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Prezenta pe Google Maps" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_key_takeaways_none_omits_section(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, {"key_takeaways": None})
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Puncte Cheie" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_sector_position_empty_dict_omits_section(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        verified_data = {"risk_score": {"sector_position": {}}}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, verified_data)
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Pozitie in Sector" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)


class TestTavilyQuotaAndDivergenceDocx:
    """A6 + A4 (2026-07-16): mesajul onest de cota Tavily epuizata + dezacordul
    FAPTIC scor 6D vs modele predictive, randate in DOCX. Combinate cu fixture-ul
    diacritic-greu (AEGRM/historical_flags) deja verificat, ca proba ca sectiunile
    noi nu strica generarea existenta."""

    def _combined_fixture(self) -> dict:
        return {
            "tavily_quota_exhausted": {"value": True, "usage": 950},
            "risk_score": {"score": "Verde", "numeric_score": 78.0},
            "predictive_scores": {
                "altman_z": {"z_score": None, "zone": "INDISPONIBIL"},
                "piotroski_f": {"f_score": 4, "max_possible": 5, "grade": "STRONG"},
                "beneish_m": {"m_score": None, "risk": "INDISPONIBIL", "available": False},
                "zmijewski_x": {"x_score": 2.4, "distress": True, "available": True},
                "distress_signals": 1,
                "summary": "Firma in zona gri (1/4 modele calculate)",
            },
            # Diacritics context, same shape as test_functioneaza_cu_rich_fields_aegrm_historical.
            "historical_flags": [
                {"type": "cesiune_parti_sociale", "label": "Cesiune părți sociale detectată",
                 "severity": "HIGH", "snippet": "Schimbare asociați — cesiune 60% părți sociale"},
            ],
        }

    def test_content_rendered_with_diacritics_context(self):
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, self._combined_fixture())
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)

            assert "Verificare Incompleta" in full_text
            assert "NU a fost efectuata" in full_text
            assert "950/1000" in full_text
            assert "Dezacord intre scorul 6D si modelele predictive" in full_text
            assert "Scor 6D: Verde (78.0)" in full_text
            assert "semnal de distres" in full_text
            assert "Cele doua metode nu concorda" in full_text
            # Fixture-ul diacritic invecinat tot randeaza corect.
            assert "Cesiune părți sociale detectată" in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)

    def test_no_divergence_when_models_agree(self):
        """Caz real (TAROM): scor Verde + Zmijewski fara distres -- fara bloc de dezacord."""
        from docx import Document

        from backend.reports.docx_generator import generate_docx

        data = self._combined_fixture()
        data["tavily_quota_exhausted"] = {}
        data["predictive_scores"]["zmijewski_x"] = {"x_score": -0.85, "distress": False, "available": True}
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            generate_docx(_make_sections(), _make_meta(), path, data)
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Verificare Incompleta" not in full_text
            assert "Dezacord intre scorul 6D" not in full_text
        finally:
            if os.path.exists(path):
                os.remove(path)
